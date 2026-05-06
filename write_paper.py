import os
from docx import Document
from docx.shared import Inches

def write_paper():
    template_path = r'C:\Users\Márcio\Desktop\Tornado_Building_Simulator\Conference_Paper_Template.docx'
    output_path = r'C:\Users\Márcio\Desktop\Tornado_Building_Simulator\Final_Paper_Marcio_Ferreira.docx'
    desktop_output_path = r'C:\Users\Márcio\Desktop\Final_Paper_Marcio_Ferreira.docx'
    images_dir = r'C:\Users\Márcio\Desktop\Tornado_Building_Simulator\extracted_images'

    # Try to open template, if it fails, create a blank one
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Could not open template due to {e}. Creating a blank document.")
        doc = Document()
        
    doc.add_heading('TORNADO STRUCTURAL RETROFIT OPTIMIZATION', 0)
    doc.add_heading('STRUCTURAL PRESERVATION IN UNREINFORCED MASONRY HISTORIC BUILDINGS', 1)
    doc.add_paragraph('Márcio Ferreira | March 2026\nAe 540 - Professor Nathan Brown')

    # 1. Motivation
    doc.add_heading('1. Motivation', level=1)
    doc.add_paragraph(
        "19th-century masonry structures are a cornerstone of urban heritage, but they possess a significant structural vulnerability: the fire-cut joist. "
        "Originally designed to prevent wall collapse during a fire by allowing floor joists to slide out, this connection offers negligible resistance to modern design wind speeds. "
        "With current ASCE 7-22 standards requiring resistance to 135 mph winds (generating roughly 2,062 lbf of uplift in our test case), these buildings face a high risk of catastrophic roof blow-offs. "
        "This creates a profound friction between preserving historic fabric and ensuring life safety. To address this, we aim to automate the search for optimal structural retrofit solutions using computational methods."
    )
    
    # Add an image if available
    img1 = os.path.join(images_dir, 'image_1_0.png')
    if os.path.exists(img1):
        doc.add_picture(img1, width=Inches(4))

    # 2. Literature Review
    doc.add_heading('2. Literature Review', level=1)
    doc.add_paragraph(
        "The evaluation of unreinforced masonry under lateral and uplift loads relies heavily on the ASCE 41-23 standard for Seismic/Wind Evaluation. "
        "Commercially available advanced connector specifications, such as those provided by Simpson Strong-Tie, offer discrete mechanical interventions for roof-to-wall anchoring. "
        "Furthermore, recent advancements by Napolitano et al. regarding Heritage Digital Twin methodologies provide a framework for continuous structural health monitoring. "
        "Current state-of-the-art reviews on wall-to-horizontal diaphragm connections in historical buildings highlight the necessity for minimally invasive yet structurally robust ties to prevent out-of-plane failure."
    )
    
    img2 = os.path.join(images_dir, 'image_3_3.png')
    if os.path.exists(img2):
        doc.add_picture(img2, width=Inches(3))

    # 3. Methodology
    doc.add_heading('3. Methodology', level=1)
    doc.add_paragraph(
        "The computational pipeline is divided into four main stages: Parametric Scripting, Genetic Heuristics, Abaqus Standard/Explicit Structural Analysis, and a Digital Interactive Dashboard. "
        "The overarching goal is to achieve optimal preservation alongside maximum wind resistance for EF-4 tornado events. "
        "Finite Element Analysis (FEA) models, built via Abaqus macros, are employed because they are accurate and computationally efficient. "
        "Monte Carlo uncertainty quantifies the associated risks, while Pareto-optimal solutions generated via Genetic Algorithms enable informed, reproducible decision-making."
    )

    # 4. Results
    doc.add_heading('4. Results', level=1)
    doc.add_heading('The Optimization Dashboard', level=2)
    doc.add_paragraph(
        "The workflow begins in the Optimization Dashboard, where the environmental demand is established. Inputting the physical parameters (roof span, joist spacing) and the tornadic wind speed computes the Design Uplift Demand. "
        "The genetic algorithm then evaluates 8 different structural retrofit strategies, offloading the heavy computational search to the Penn State Roar Supercomputer."
    )
    
    img3 = os.path.join(images_dir, 'image_10_13.png')
    if os.path.exists(img3):
        doc.add_picture(img3, width=Inches(5))

    doc.add_heading('Pareto Front Visual Report', level=2)
    doc.add_paragraph(
        "The React-based Visual Report presents the results through a Pareto Frontier. This highlights the trade-off between Safety (maximizing uplift capacity) and Historic Impact (minimizing invasiveness). "
        "Any point on this curve represents a mathematically non-dominated strategy. Selecting a 'Winning Strategy' packages the exact geometric and material parameters into the 3D Finite Element Simulator."
    )

    doc.add_heading('3D Building Simulator & Abaqus Export', level=2)
    doc.add_paragraph(
        "The 3D Building Simulator generates the full building model using the selected hardware parameters. The macro-geometry (story height, wall thickness, window placement) and Abaqus solver settings (Concrete Damage Plasticity viscosity, mass scaling) are configured natively. "
        "Upon execution, a Python engine dynamically writes and runs an Abaqus Macro. The dashboard then projects failure mappings, such as Tensile Damage (DAMAGET) and Scalar Stiffness Degradation (SDEG)."
    )

    # 5. Discussion
    doc.add_heading('5. Discussion', level=1)
    doc.add_paragraph(
        "The integration of parametric modeling with advanced heuristic search provides a highly defensible approach to heritage retrofitting. "
        "By avoiding blind guesses of hardware sizes in Abaqus, the optimization ensures that the selected retrofit does not unnecessarily destroy 19th-century brickwork. "
        "Furthermore, bypassing the UI to open the fully rendered .odb database live allows for advanced forensic investigations, offering profound insights into localized out-of-plane wall failures and stress concentrations around openings."
    )

    # 6. Summary of contributions and future work
    doc.add_heading('6. Summary of Contributions and Future Work', level=1)
    doc.add_paragraph(
        "The primary contribution of this work is the end-to-end automation of retrofit optimization for historic unreinforced masonry, bridging the gap between theoretical FEA simulations and real-world structural management. "
        "Future research directions include:\n"
        "• Expansion to Multi-Story Meso-Scale Discretization: Expanding the explicit brick-by-brick mesh generation to full-scale, multi-story buildings to analyze global overturning versus localized failure.\n"
        "• FSI via CFD Coupling: Coupling the structural Abaqus model with Computational Fluid Dynamics (CFD) to capture realistic turbulent vortex behaviors and dynamic wind shedding.\n"
        "• Machine Learning Surrogate Models: Training an advanced Deep Learning surrogate model using massive synthetic datasets from the ROAR supercomputer to instantly predict SDEG and collapse probability.\n"
        "• Structural Digital Twins: Linking predictive damage thresholds with real-world IoT sensor data (e.g., tiltmeters, strain gauges) to create a living virtual replica capable of proactive disaster readiness and real-time health monitoring."
    )

    doc.save(output_path)
    try:
        doc.save(desktop_output_path)
    except Exception as e:
        print(f"Failed to save to desktop: {e}")
        
    print(f"Paper saved successfully to {output_path} and {desktop_output_path}")

if __name__ == "__main__":
    write_paper()
